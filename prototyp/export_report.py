import os
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def add_table_from_csv(doc, csv_path, title):
    """Vloží tabuľku z CSV do dokumentu so zarámovaním buniek."""
    df = pd.read_csv(csv_path, sep=";")
    doc.add_heading(title, level=2)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    # Header
    hdr_cells = table.rows[0].cells
    for i, column_name in enumerate(df.columns):
        hdr_cells[i].text = column_name

    # Data
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    doc.add_paragraph()

def add_image(doc, image_path, title):
    """Vloží obrázok do dokumentu."""
    doc.add_heading(title, level=2)
    doc.add_picture(image_path, width=Inches(5))
    doc.add_paragraph()

def generate_report():
    # === Cesty k súborom ===
    cleaned_dataset_path = "sem_SSBU/SSBU25_dataset_cleaned.csv"
    grafy_dir = "grafy"
    tabulky_dir = "tabulky"

    # === Načítaj dataset (kvôli úvodu) ===
    df = pd.read_csv(cleaned_dataset_path, sep=";", encoding="utf-8-sig")

    # === Vytvor dokument ===
    doc = Document()

    # === Titulná strana ===
    doc.add_heading("Analýza HFE génu a genetickej predispozície na hemochromatózu", 0)
    doc.add_paragraph("Generovaný report").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_page_break()

    # === Obsah ===
    doc.add_heading('Obsah', level=1)
    doc.add_paragraph("1. Základné informácie o datasete")
    doc.add_paragraph("2. Hardy-Weinbergova rovnováha")
    doc.add_paragraph("3. Percentuálne zastúpenie genotypov a prenášači")
    doc.add_paragraph("4. Súvislosť HFE mutácií s pečeňovými diagnózami")
    doc.add_paragraph("5. Grafy rozdelenia genotypov a ich vzťah k atribútom")
    doc.add_paragraph("6. Analýza diagnóz podľa MKCH-10")
    doc.add_page_break()

    # === 1. Základné informácie ===
    doc.add_heading('1. Základné informácie o datasete', level=1)
    doc.add_paragraph(f"Počet riadkov: {len(df)}")
    doc.add_paragraph(f"Počet stĺpcov: {len(df.columns)}")
    doc.add_paragraph("Stĺpce:")
    for col in df.columns:
        doc.add_paragraph(f"- {col}", style='List Bullet')
    doc.add_page_break()

    # === 2. Hardy-Weinbergova rovnováha ===
    doc.add_heading('2. Hardy-Weinbergova rovnováha', level=1)
    hwe_path = os.path.join(tabulky_dir, "hardy_weinberg_test.csv")
    if os.path.exists(hwe_path):
        add_table_from_csv(doc, hwe_path, "Výsledky Hardy-Weinbergovho testu")
    else:
        doc.add_paragraph("Hardy-Weinbergove výsledky neboli nájdené.")
    doc.add_page_break()

    # === 3. Percentuálne zastúpenie genotypov a prenášači ===
    doc.add_heading('3. Percentá genotypov a prenášači', level=1)
    percenta_path = os.path.join(tabulky_dir, "percenta_genotypov.csv")
    pren_path = os.path.join(tabulky_dir, "prenasic_predispozicia.txt")

    if os.path.exists(percenta_path):
        add_table_from_csv(doc, percenta_path, "Percentá genotypov")
    if os.path.exists(pren_path):
        with open(pren_path, "r", encoding="utf-8") as f:
            doc.add_paragraph(f.read())
    doc.add_page_break()

    # === 4. Súvislosť HFE mutácií s pečeňovými diagnózami ===
    doc.add_heading('4. Súvislosť HFE mutácií s pečeňovými diagnózami', level=1)
    suvislost_path = os.path.join(tabulky_dir, "vysledky_suvislosti.csv")
    if os.path.exists(suvislost_path):
        add_table_from_csv(doc, suvislost_path, "Výsledky súvislostí")
    for mut in ["H63D", "S65C", "C282Y"]:
        graf_path = os.path.join(grafy_dir, f"graf_suvislost_{mut}.png")
        if os.path.exists(graf_path):
            add_image(doc, graf_path, f"Graf súvislosti pre mutáciu {mut}")
    doc.add_page_break()

    # === 5. Grafy rozdelenia genotypov ===
    doc.add_heading('5. Grafy rozdelenia genotypov', level=1)
    for mut in ["H63D", "S65C", "C282Y"]:
        for typ in ["rozdelenie", "vek_vs_genotyp", "pohlavie_vs_genotyp", "pecen_vs_genotyp"]:
            graf_path = os.path.join(grafy_dir, f"{typ}_{mut}.png")
            if os.path.exists(graf_path):
                add_image(doc, graf_path, f"{typ.replace('_', ' ').capitalize()} – {mut}")
    doc.add_page_break()

    # === 6. Analýza diagnóz podľa MKCH-10 ===
    doc.add_heading('6. Analýza diagnóz podľa MKCH-10', level=1)
    mkch_path = os.path.join(tabulky_dir, "vyvoj_skupin_diagnoz.csv")
    if os.path.exists(mkch_path):
        add_table_from_csv(doc, mkch_path, "Vývoj diagnóz v čase")
    mkch_graf_path = os.path.join(grafy_dir, "vyvoj_skupin_diagnoz.png")
    if os.path.exists(mkch_graf_path):
        add_image(doc, mkch_graf_path, "Graf vývoja skupín diagnóz")
    zastarale_path = os.path.join(tabulky_dir, "zastarale_kody_diagnoz.csv")
    if os.path.exists(zastarale_path):
        add_table_from_csv(doc, zastarale_path, "Zastarané kódy diagnóz")
    doc.add_page_break()

    # === Ulož dokument ===
    output_path = "HFE_gene_analysis_report.docx"
    doc.save(output_path)

    return output_path
